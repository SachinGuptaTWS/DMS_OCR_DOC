import time
import os
import traceback
from docx import Document

def execute_docx_ocr(filepath):
    """
    Extracts content from .docx files using python-docx.
    - Preserves Paragraphs
    - Converts Tables to Markdown for the frontend renderer
    """
    start_time = time.time()
    extracted_text = ""
    slides_data = [] # Reusing 'slides' structure for pages/sections
    status = "Success"
    
    print(f"   [+] Analyzing Word Doc: {os.path.basename(filepath)}")

    try:
        doc = Document(filepath)
        
        # We will group content into a single "Page" for simplicity, 
        # or split by sections if needed. Here we do one continuous flow.
        items = []
        full_text_buffer = []

        # Iterate through the document's body elements in order
        # Note: python-docx separates paragraphs and tables. 
        # For a simple implementation, we iterate paragraphs then tables, 
        # or we try to reconstruct order (complex). 
        # Strategy: Iterate paragraphs for text, tables for structure.
        
        # 1. Extract Text Paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                full_text_buffer.append(text)
                items.append({
                    "type": "text_block",
                    "text": text
                })

        # 2. Extract Tables
        # We append them to the items list (rendering order might be separate from paragraphs in python-docx API)
        for i, table in enumerate(doc.tables):
            md_table = ""
            rows_data = []
            
            for row in table.rows:
                row_cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                rows_data.append(row_cells)
            
            if rows_data:
                # Header
                header = rows_data[0]
                md_table += "| " + " | ".join(header) + " |\n"
                md_table += "| " + " | ".join(["---"] * len(header)) + " |\n"
                
                # Body
                for r in rows_data[1:]:
                    md_table += "| " + " | ".join(r) + " |\n"
                
                items.append({
                    "type": "table_structure",
                    "text": md_table,
                    "conf": 1.0,
                    "source": f"Native Word Table #{i+1}"
                })
                
                # Add table text to full text buffer for NLP analysis
                full_text_buffer.append(f"\n[Table Data: {', '.join(header)}...]\n")

        extracted_text = "\n".join(full_text_buffer)

        slides_data.append({
            "slide_number": 1,
            "items": items,
            "accuracy_score": 100
        })

    except Exception as e:
        status = f"Error: {str(e)[:50]}"
        traceback.print_exc()
        extracted_text = ""

    duration = time.time() - start_time
    
    return duration, {
        "total_accuracy": 100.0, 
        "total_slides": 1, 
        "extracted_text": extracted_text, 
        "per_slide_metrics": slides_data, 
        "ner_entities": []
    }, "Native Python-Docx", status